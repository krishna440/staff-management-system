from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MARKDOWN = ROOT / "PROJECT_REPORT.md"
OUTPUT = ROOT / "Staff_Management_System_Project_Report.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_width(table, width_dxa=9360):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False


def set_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr_text, fld_end])


def add_toc_placeholder(doc):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    run = p.add_run("Table of Contents")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(31, 78, 121)
    doc.add_paragraph(
        "Update this field in Microsoft Word if an automatic table of contents is required. "
        "The report sections below follow the same sequence."
    )


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in [
        ("Title", 22, RGBColor(31, 78, 121)),
        ("Heading 1", 16, RGBColor(31, 78, 121)),
        ("Heading 2", 14, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 78, 121)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = "Staff Management System Project Report"
    paragraph.runs[0].font.size = Pt(9)
    paragraph.runs[0].font.color.rgb = RGBColor(99, 99, 99)
    paragraph.paragraph_format.space_after = Pt(4)

    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D9E2F3")
    border.append(bottom)
    paragraph._p.get_or_add_pPr().append(border)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    set_page_number(paragraph)


def parse_inline_bold(paragraph, text):
    parts = text.split("**")
    for index, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.bold = index % 2 == 1


def add_markdown_table(doc, rows):
    headers = [cell.strip() for cell in rows[0].strip("|").split("|")]
    body_rows = rows[2:]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table)

    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        set_cell_shading(header_cells[idx], "EAF2F8")
        set_cell_margins(header_cells[idx])
        header_cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in header_cells[idx].paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.name = "Arial"
            p.runs[0].font.size = Pt(9)
    set_repeat_table_header(table.rows[0])

    for row in body_rows:
        values = [cell.strip() for cell in row.strip("|").split("|")]
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = value
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[idx].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)

    doc.add_paragraph()


def build_doc():
    doc = Document()
    style_document(doc)
    add_header_footer(doc)

    lines = MARKDOWN.read_text(encoding="utf-8").splitlines()
    table_buffer = []
    in_title_block = True

    for line in lines:
        stripped = line.strip()

        if table_buffer and (not stripped.startswith("|")):
            add_markdown_table(doc, table_buffer)
            table_buffer = []

        if not stripped:
            continue

        if stripped.startswith("|"):
            table_buffer.append(stripped)
            continue

        if stripped.startswith("# "):
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            parse_inline_bold(p, stripped[2:])
            doc.add_paragraph()
            continue

        if stripped.startswith("## "):
            title = stripped[3:]
            if title == "Table of Contents":
                add_toc_placeholder(doc)
            else:
                doc.add_heading(title, level=1)
            in_title_block = False
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            parse_inline_bold(p, stripped[2:])
            continue

        if stripped[:3].isdigit() and ". " in stripped[:5]:
            p = doc.add_paragraph(style="List Number")
            parse_inline_bold(p, stripped.split(". ", 1)[1])
            continue

        p = doc.add_paragraph()
        if in_title_block:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        parse_inline_bold(p, stripped)

    if table_buffer:
        add_markdown_table(doc, table_buffer)

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_doc())
